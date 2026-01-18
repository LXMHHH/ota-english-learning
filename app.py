#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OTA English Learning Web App V2.0
OTA测试工程师英语学习工具 - 增强版

新增功能:
- 200+词汇
- 100+短语
- 文档导入功能 - 可上传英文文档自动提取单词和句子

运行方式: streamlit run ota_english_app_v2.py
"""

import streamlit as st
import streamlit.components.v1 as components
import random
import json
import re
from datetime import datetime, date
from pathlib import Path

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="OTA英语学习",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ==================== PWA支持 ====================
st.markdown("""
<head>
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
    <meta name="apple-mobile-web-app-title" content="OTA英语">
    <meta name="mobile-web-app-capable" content="yes">
    <meta name="viewport" content="width=device-width, initial-scale=1, maximum-scale=1, user-scalable=no">
</head>
""", unsafe_allow_html=True)

# ==================== 样式 ====================
st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    .stButton > button {width: 100%; padding: 12px; font-size: 16px; border-radius: 10px;}
    .word-card {background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 15px; margin: 10px 0;}
    .phrase-card {background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%); color: white; padding: 15px; border-radius: 12px; margin: 8px 0;}
    .flashcard {background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); color: white; padding: 30px; border-radius: 20px; text-align: center; min-height: 180px; margin: 15px 0;}
    .speak-btn {
        background: #4CAF50;
        color: white;
        border: none;
        padding: 8px 16px;
        border-radius: 8px;
        cursor: pointer;
        font-size: 14px;
        margin: 5px 2px;
        display: inline-block;
        transition: all 0.3s;
    }
    .speak-btn:hover {
        background: #45a049;
        transform: scale(1.05);
    }
    .speak-btn:active {
        transform: scale(0.95);
    }
</style>
""", unsafe_allow_html=True)

# ==================== 扩展词汇数据 (200+) ====================
OTA_VOCABULARY = [
    # ECU Communication (30)
    {"id": "v001", "english": "ECU", "chinese": "电子控制单元", "phonetic": "/ˌiː siː ˈjuː/", "example": "The ECU controls the vehicle's engine.", "category": "ECU通信"},
    {"id": "v002", "english": "firmware", "chinese": "固件", "phonetic": "/ˈfɜːmweə/", "example": "Update the ECU firmware via OTA.", "category": "ECU通信"},
    {"id": "v003", "english": "bootloader", "chinese": "引导程序", "phonetic": "/ˈbuːtˌləʊdə/", "example": "The bootloader initializes the system.", "category": "ECU通信"},
    {"id": "v004", "english": "flash memory", "chinese": "闪存", "phonetic": "/flæʃ ˈmeməri/", "example": "Firmware is stored in flash memory.", "category": "ECU通信"},
    {"id": "v005", "english": "CAN bus", "chinese": "CAN总线", "phonetic": "/kæn bʌs/", "example": "ECUs communicate through CAN bus.", "category": "ECU通信"},
    {"id": "v006", "english": "diagnostic", "chinese": "诊断", "phonetic": "/ˌdaɪəɡˈnɒstɪk/", "example": "Run diagnostic tests on ECU.", "category": "ECU通信"},
    {"id": "v007", "english": "UDS", "chinese": "统一诊断服务", "phonetic": "/ˌjuː diː ˈes/", "example": "UDS protocol for ECU diagnostics.", "category": "ECU通信"},
    {"id": "v008", "english": "DoIP", "chinese": "IP诊断", "phonetic": "/dəʊ aɪ ˈpiː/", "example": "DoIP enables diagnostics over IP.", "category": "ECU通信"},
    {"id": "v009", "english": "DID", "chinese": "数据标识符", "phonetic": "/diː aɪ ˈdiː/", "example": "Read version using DID F189.", "category": "ECU通信"},
    {"id": "v010", "english": "session", "chinese": "会话", "phonetic": "/ˈseʃən/", "example": "Switch to extended session.", "category": "ECU通信"},
    {"id": "v011", "english": "security access", "chinese": "安全访问", "phonetic": "/sɪˈkjʊərəti/", "example": "Security access required.", "category": "ECU通信"},
    {"id": "v012", "english": "seed and key", "chinese": "种子密钥", "phonetic": "/siːd ænd kiː/", "example": "Seed and key algorithm.", "category": "ECU通信"},
    {"id": "v013", "english": "NRC", "chinese": "否定响应码", "phonetic": "/en ɑː siː/", "example": "NRC 0x22 conditions not correct.", "category": "ECU通信"},
    {"id": "v014", "english": "positive response", "chinese": "肯定响应", "phonetic": "/ˈpɒzətɪv/", "example": "Received positive response.", "category": "ECU通信"},
    {"id": "v015", "english": "timeout", "chinese": "超时", "phonetic": "/ˈtaɪmaʊt/", "example": "Request timed out.", "category": "ECU通信"},
    {"id": "v016", "english": "VCU", "chinese": "整车控制器", "phonetic": "/viː siː juː/", "example": "VCU controls vehicle functions.", "category": "ECU通信"},
    {"id": "v017", "english": "BCM", "chinese": "车身控制模块", "phonetic": "/biː siː em/", "example": "BCM manages body electronics.", "category": "ECU通信"},
    {"id": "v018", "english": "BMS", "chinese": "电池管理系统", "phonetic": "/biː em es/", "example": "BMS monitors battery status.", "category": "ECU通信"},
    {"id": "v019", "english": "MCU", "chinese": "电机控制器", "phonetic": "/em siː juː/", "example": "MCU controls the motor.", "category": "ECU通信"},
    {"id": "v020", "english": "IVI", "chinese": "车载娱乐系统", "phonetic": "/aɪ viː aɪ/", "example": "IVI provides entertainment.", "category": "ECU通信"},
    {"id": "v021", "english": "cluster", "chinese": "仪表盘", "phonetic": "/ˈklʌstə/", "example": "Update cluster display.", "category": "ECU通信"},
    {"id": "v022", "english": "gateway", "chinese": "网关", "phonetic": "/ˈɡeɪtweɪ/", "example": "Gateway routes messages.", "category": "ECU通信"},
    {"id": "v023", "english": "ADAS", "chinese": "驾驶辅助系统", "phonetic": "/ˈeɪdæs/", "example": "ADAS requires updates.", "category": "ECU通信"},
    {"id": "v024", "english": "calibration", "chinese": "标定", "phonetic": "/ˌkælɪˈbreɪʃən/", "example": "Update calibration data.", "category": "ECU通信"},
    {"id": "v025", "english": "parameter", "chinese": "参数", "phonetic": "/pəˈræmɪtə/", "example": "Modify ECU parameters.", "category": "ECU通信"},
    {"id": "v026", "english": "tester present", "chinese": "测试仪在线", "phonetic": "/ˈtestə/", "example": "Send tester present.", "category": "ECU通信"},
    {"id": "v027", "english": "routine", "chinese": "例程", "phonetic": "/ruːˈtiːn/", "example": "Execute diagnostic routine.", "category": "ECU通信"},
    {"id": "v028", "english": "memory address", "chinese": "内存地址", "phonetic": "/ˈmeməri/", "example": "Specify memory address.", "category": "ECU通信"},
    {"id": "v029", "english": "data length", "chinese": "数据长度", "phonetic": "/ˈdeɪtə leŋθ/", "example": "Check data length.", "category": "ECU通信"},
    {"id": "v030", "english": "block transfer", "chinese": "块传输", "phonetic": "/blɒk/", "example": "Use block transfer mode.", "category": "ECU通信"},

    # OTA Process (40)
    {"id": "v031", "english": "OTA", "chinese": "空中下载", "phonetic": "/əʊ tiː eɪ/", "example": "OTA update allows remote updates.", "category": "OTA流程"},
    {"id": "v032", "english": "FOTA", "chinese": "固件空中下载", "phonetic": "/fəʊtə/", "example": "FOTA for firmware updates.", "category": "OTA流程"},
    {"id": "v033", "english": "SOTA", "chinese": "软件空中下载", "phonetic": "/səʊtə/", "example": "SOTA updates application.", "category": "OTA流程"},
    {"id": "v034", "english": "download", "chinese": "下载", "phonetic": "/ˈdaʊnləʊd/", "example": "Download update package.", "category": "OTA流程"},
    {"id": "v035", "english": "install", "chinese": "安装", "phonetic": "/ɪnˈstɔːl/", "example": "Install the update.", "category": "OTA流程"},
    {"id": "v036", "english": "activate", "chinese": "激活", "phonetic": "/ˈæktɪveɪt/", "example": "Activate new version.", "category": "OTA流程"},
    {"id": "v037", "english": "rollback", "chinese": "回滚", "phonetic": "/ˈrəʊlbæk/", "example": "Rollback if update fails.", "category": "OTA流程"},
    {"id": "v038", "english": "delta update", "chinese": "差分升级", "phonetic": "/ˈdeltə/", "example": "Delta update reduces size.", "category": "OTA流程"},
    {"id": "v039", "english": "full update", "chinese": "全量升级", "phonetic": "/fʊl/", "example": "Full update replaces all.", "category": "OTA流程"},
    {"id": "v040", "english": "package", "chinese": "升级包", "phonetic": "/ˈpækɪdʒ/", "example": "Update package is 500MB.", "category": "OTA流程"},
    {"id": "v041", "english": "manifest", "chinese": "清单文件", "phonetic": "/ˈmænɪfest/", "example": "Manifest contains metadata.", "category": "OTA流程"},
    {"id": "v042", "english": "checksum", "chinese": "校验和", "phonetic": "/ˈtʃeksʌm/", "example": "Verify package checksum.", "category": "OTA流程"},
    {"id": "v043", "english": "signature", "chinese": "签名", "phonetic": "/ˈsɪɡnətʃə/", "example": "Verify digital signature.", "category": "OTA流程"},
    {"id": "v044", "english": "campaign", "chinese": "升级活动", "phonetic": "/kæmˈpeɪn/", "example": "Launch OTA campaign.", "category": "OTA流程"},
    {"id": "v045", "english": "progress", "chinese": "进度", "phonetic": "/ˈprəʊɡres/", "example": "Download progress 80%.", "category": "OTA流程"},
    {"id": "v046", "english": "server", "chinese": "服务器", "phonetic": "/ˈsɜːvə/", "example": "Connect to OTA server.", "category": "OTA流程"},
    {"id": "v047", "english": "client", "chinese": "客户端", "phonetic": "/ˈklaɪənt/", "example": "Vehicle is OTA client.", "category": "OTA流程"},
    {"id": "v048", "english": "API", "chinese": "接口", "phonetic": "/eɪ piː aɪ/", "example": "Call the OTA API.", "category": "OTA流程"},
    {"id": "v049", "english": "MQTT", "chinese": "消息队列", "phonetic": "/em kjuː tiː tiː/", "example": "Use MQTT for messaging.", "category": "OTA流程"},
    {"id": "v050", "english": "certificate", "chinese": "证书", "phonetic": "/səˈtɪfɪkət/", "example": "Verify SSL certificate.", "category": "OTA流程"},
    {"id": "v051", "english": "encryption", "chinese": "加密", "phonetic": "/ɪnˈkrɪpʃən/", "example": "Data encryption required.", "category": "OTA流程"},
    {"id": "v052", "english": "authentication", "chinese": "认证", "phonetic": "/ɔːˌθentɪˈkeɪʃən/", "example": "Vehicle authentication.", "category": "OTA流程"},
    {"id": "v053", "english": "hash", "chinese": "哈希值", "phonetic": "/hæʃ/", "example": "Calculate file hash.", "category": "OTA流程"},
    {"id": "v054", "english": "compress", "chinese": "压缩", "phonetic": "/kəmˈpres/", "example": "Compress the package.", "category": "OTA流程"},
    {"id": "v055", "english": "decompress", "chinese": "解压", "phonetic": "/diːkəmˈpres/", "example": "Decompress before install.", "category": "OTA流程"},
    {"id": "v056", "english": "partition", "chinese": "分区", "phonetic": "/pɑːˈtɪʃən/", "example": "Update system partition.", "category": "OTA流程"},
    {"id": "v057", "english": "A/B update", "chinese": "A/B升级", "phonetic": "/eɪ biː/", "example": "Use A/B update method.", "category": "OTA流程"},
    {"id": "v058", "english": "slot", "chinese": "槽位", "phonetic": "/slɒt/", "example": "Switch to slot B.", "category": "OTA流程"},
    {"id": "v059", "english": "backup", "chinese": "备份", "phonetic": "/ˈbækʌp/", "example": "Backup before update.", "category": "OTA流程"},
    {"id": "v060", "english": "restore", "chinese": "恢复", "phonetic": "/rɪˈstɔː/", "example": "Restore from backup.", "category": "OTA流程"},
    {"id": "v061", "english": "retry", "chinese": "重试", "phonetic": "/riːˈtraɪ/", "example": "Retry the download.", "category": "OTA流程"},
    {"id": "v062", "english": "resume", "chinese": "断点续传", "phonetic": "/rɪˈzjuːm/", "example": "Resume download.", "category": "OTA流程"},
    {"id": "v063", "english": "prerequisite", "chinese": "前置条件", "phonetic": "/priːˈrekwɪzɪt/", "example": "Check prerequisites.", "category": "OTA流程"},
    {"id": "v064", "english": "dependency", "chinese": "依赖", "phonetic": "/dɪˈpendənsi/", "example": "Check dependencies.", "category": "OTA流程"},
    {"id": "v065", "english": "compatible", "chinese": "兼容", "phonetic": "/kəmˈpætəbəl/", "example": "Version compatible.", "category": "OTA流程"},
    {"id": "v066", "english": "trigger", "chinese": "触发", "phonetic": "/ˈtrɪɡə/", "example": "Trigger the update.", "category": "OTA流程"},
    {"id": "v067", "english": "notification", "chinese": "通知", "phonetic": "/ˌnəʊtɪfɪˈkeɪʃən/", "example": "Send notification.", "category": "OTA流程"},
    {"id": "v068", "english": "consent", "chinese": "同意", "phonetic": "/kənˈsent/", "example": "User consent required.", "category": "OTA流程"},
    {"id": "v069", "english": "queue", "chinese": "队列", "phonetic": "/kjuː/", "example": "Add to update queue.", "category": "OTA流程"},
    {"id": "v070", "english": "batch", "chinese": "批次", "phonetic": "/bætʃ/", "example": "Update in batches.", "category": "OTA流程"},
    
    # Vehicle Signals (25)
    {"id": "v071", "english": "ignition", "chinese": "点火", "phonetic": "/ɪɡˈnɪʃən/", "example": "Turn off ignition.", "category": "车辆信号"},
    {"id": "v072", "english": "voltage", "chinese": "电压", "phonetic": "/ˈvəʊltɪdʒ/", "example": "Battery voltage above 12V.", "category": "车辆信号"},
    {"id": "v073", "english": "speed", "chinese": "车速", "phonetic": "/spiːd/", "example": "Vehicle speed must be zero.", "category": "车辆信号"},
    {"id": "v074", "english": "gear", "chinese": "档位", "phonetic": "/ɡɪə/", "example": "Put gear in Park.", "category": "车辆信号"},
    {"id": "v075", "english": "handbrake", "chinese": "手刹", "phonetic": "/ˈhænbreɪk/", "example": "Engage handbrake.", "category": "车辆信号"},
    {"id": "v076", "english": "battery", "chinese": "电池", "phonetic": "/ˈbætəri/", "example": "Connect external battery.", "category": "车辆信号"},
    {"id": "v077", "english": "T-Box", "chinese": "车载终端", "phonetic": "/tiː bɒks/", "example": "T-Box handles OTA.", "category": "车辆信号"},
    {"id": "v078", "english": "VIN", "chinese": "车辆识别码", "phonetic": "/vɪn/", "example": "Verify VIN.", "category": "车辆信号"},
    {"id": "v079", "english": "network", "chinese": "网络", "phonetic": "/ˈnetwɜːk/", "example": "Check network.", "category": "车辆信号"},
    {"id": "v080", "english": "signal strength", "chinese": "信号强度", "phonetic": "/ˈsɪɡnəl/", "example": "Signal strength weak.", "category": "车辆信号"},
    {"id": "v081", "english": "engine", "chinese": "发动机", "phonetic": "/ˈendʒɪn/", "example": "Engine must be off.", "category": "车辆信号"},
    {"id": "v082", "english": "door", "chinese": "车门", "phonetic": "/dɔː/", "example": "All doors closed.", "category": "车辆信号"},
    {"id": "v083", "english": "temperature", "chinese": "温度", "phonetic": "/ˈtemprətʃə/", "example": "ECU temperature normal.", "category": "车辆信号"},
    {"id": "v084", "english": "odometer", "chinese": "里程表", "phonetic": "/əʊˈdɒmɪtə/", "example": "Record odometer reading.", "category": "车辆信号"},
    {"id": "v085", "english": "mileage", "chinese": "里程", "phonetic": "/ˈmaɪlɪdʒ/", "example": "Current mileage 50000km.", "category": "车辆信号"},
    {"id": "v086", "english": "charging", "chinese": "充电", "phonetic": "/ˈtʃɑːdʒɪŋ/", "example": "Vehicle is charging.", "category": "车辆信号"},
    {"id": "v087", "english": "SOC", "chinese": "电量百分比", "phonetic": "/es əʊ siː/", "example": "SOC above 30%.", "category": "车辆信号"},
    {"id": "v088", "english": "driving mode", "chinese": "驾驶模式", "phonetic": "/ˈdraɪvɪŋ/", "example": "Switch driving mode.", "category": "车辆信号"},
    {"id": "v089", "english": "parking", "chinese": "驻车", "phonetic": "/ˈpɑːkɪŋ/", "example": "Vehicle in parking.", "category": "车辆信号"},
    {"id": "v090", "english": "ready", "chinese": "就绪", "phonetic": "/ˈredi/", "example": "Vehicle ready state.", "category": "车辆信号"},
    {"id": "v091", "english": "standby", "chinese": "待机", "phonetic": "/ˈstændbaɪ/", "example": "Enter standby mode.", "category": "车辆信号"},
    {"id": "v092", "english": "wake up", "chinese": "唤醒", "phonetic": "/weɪk ʌp/", "example": "Wake up the ECU.", "category": "车辆信号"},
    {"id": "v093", "english": "sleep", "chinese": "休眠", "phonetic": "/sliːp/", "example": "ECU enters sleep.", "category": "车辆信号"},
    {"id": "v094", "english": "power supply", "chinese": "电源", "phonetic": "/ˈpaʊə/", "example": "Check power supply.", "category": "车辆信号"},
    {"id": "v095", "english": "ground", "chinese": "接地", "phonetic": "/ɡraʊnd/", "example": "Check ground connection.", "category": "车辆信号"},

    # Diagnostic Protocols (25)
    {"id": "v096", "english": "protocol", "chinese": "协议", "phonetic": "/ˈprəʊtəkɒl/", "example": "Use UDS protocol.", "category": "诊断协议"},
    {"id": "v097", "english": "service", "chinese": "服务", "phonetic": "/ˈsɜːvɪs/", "example": "Service 0x34 request download.", "category": "诊断协议"},
    {"id": "v098", "english": "request", "chinese": "请求", "phonetic": "/rɪˈkwest/", "example": "Send diagnostic request.", "category": "诊断协议"},
    {"id": "v099", "english": "response", "chinese": "响应", "phonetic": "/rɪˈspɒns/", "example": "Wait for ECU response.", "category": "诊断协议"},
    {"id": "v100", "english": "transfer", "chinese": "传输", "phonetic": "/ˈtrænsfɜː/", "example": "Transfer data to ECU.", "category": "诊断协议"},
    {"id": "v101", "english": "erase", "chinese": "擦除", "phonetic": "/ɪˈreɪz/", "example": "Erase flash memory.", "category": "诊断协议"},
    {"id": "v102", "english": "write", "chinese": "写入", "phonetic": "/raɪt/", "example": "Write firmware to flash.", "category": "诊断协议"},
    {"id": "v103", "english": "read", "chinese": "读取", "phonetic": "/riːd/", "example": "Read software version.", "category": "诊断协议"},
    {"id": "v104", "english": "verify", "chinese": "校验", "phonetic": "/ˈverɪfaɪ/", "example": "Verify written data.", "category": "诊断协议"},
    {"id": "v105", "english": "reset", "chinese": "复位", "phonetic": "/riːˈset/", "example": "Reset ECU after update.", "category": "诊断协议"},
    {"id": "v106", "english": "hard reset", "chinese": "硬复位", "phonetic": "/hɑːd/", "example": "Perform hard reset.", "category": "诊断协议"},
    {"id": "v107", "english": "soft reset", "chinese": "软复位", "phonetic": "/sɒft/", "example": "Perform soft reset.", "category": "诊断协议"},
    {"id": "v108", "english": "pending", "chinese": "等待中", "phonetic": "/ˈpendɪŋ/", "example": "Response pending.", "category": "诊断协议"},
    {"id": "v109", "english": "busy", "chinese": "忙碌", "phonetic": "/ˈbɪzi/", "example": "ECU is busy.", "category": "诊断协议"},
    {"id": "v110", "english": "suppress", "chinese": "抑制", "phonetic": "/səˈpres/", "example": "Suppress positive response.", "category": "诊断协议"},
    {"id": "v111", "english": "sequence", "chinese": "序列", "phonetic": "/ˈsiːkwəns/", "example": "Block sequence number.", "category": "诊断协议"},
    {"id": "v112", "english": "address", "chinese": "地址", "phonetic": "/əˈdres/", "example": "Memory address for flash.", "category": "诊断协议"},
    {"id": "v113", "english": "length", "chinese": "长度", "phonetic": "/leŋθ/", "example": "Data length 4096 bytes.", "category": "诊断协议"},
    {"id": "v114", "english": "CRC", "chinese": "循环冗余校验", "phonetic": "/siː ɑː siː/", "example": "Calculate CRC.", "category": "诊断协议"},
    {"id": "v115", "english": "frame", "chinese": "帧", "phonetic": "/freɪm/", "example": "Send CAN frame.", "category": "诊断协议"},
    {"id": "v116", "english": "payload", "chinese": "有效载荷", "phonetic": "/ˈpeɪləʊd/", "example": "Check payload data.", "category": "诊断协议"},
    {"id": "v117", "english": "header", "chinese": "头部", "phonetic": "/ˈhedə/", "example": "Parse message header.", "category": "诊断协议"},
    {"id": "v118", "english": "acknowledge", "chinese": "确认", "phonetic": "/əkˈnɒlɪdʒ/", "example": "Send acknowledge.", "category": "诊断协议"},
    {"id": "v119", "english": "handshake", "chinese": "握手", "phonetic": "/ˈhændʃeɪk/", "example": "Complete handshake.", "category": "诊断协议"},
    {"id": "v120", "english": "baud rate", "chinese": "波特率", "phonetic": "/bɔːd reɪt/", "example": "Set baud rate 500kbps.", "category": "诊断协议"},
    
    # Test Cases (35)
    {"id": "v121", "english": "test case", "chinese": "测试用例", "phonetic": "/test keɪs/", "example": "Write test cases for OTA.", "category": "测试用例"},
    {"id": "v122", "english": "test plan", "chinese": "测试计划", "phonetic": "/test plæn/", "example": "Review the test plan.", "category": "测试用例"},
    {"id": "v123", "english": "test report", "chinese": "测试报告", "phonetic": "/test rɪˈpɔːt/", "example": "Submit test report.", "category": "测试用例"},
    {"id": "v124", "english": "pass", "chinese": "通过", "phonetic": "/pɑːs/", "example": "Test case passed.", "category": "测试用例"},
    {"id": "v125", "english": "fail", "chinese": "失败", "phonetic": "/feɪl/", "example": "Test case failed.", "category": "测试用例"},
    {"id": "v126", "english": "block", "chinese": "阻塞", "phonetic": "/blɒk/", "example": "Test is blocked.", "category": "测试用例"},
    {"id": "v127", "english": "skip", "chinese": "跳过", "phonetic": "/skɪp/", "example": "Skip this test case.", "category": "测试用例"},
    {"id": "v128", "english": "precondition", "chinese": "前置条件", "phonetic": "/priːkənˈdɪʃən/", "example": "Check preconditions.", "category": "测试用例"},
    {"id": "v129", "english": "expected result", "chinese": "预期结果", "phonetic": "/ɪkˈspektɪd/", "example": "Compare expected result.", "category": "测试用例"},
    {"id": "v130", "english": "actual result", "chinese": "实际结果", "phonetic": "/ˈæktʃuəl/", "example": "Record actual result.", "category": "测试用例"},
    {"id": "v131", "english": "regression", "chinese": "回归测试", "phonetic": "/rɪˈɡreʃən/", "example": "Run regression tests.", "category": "测试用例"},
    {"id": "v132", "english": "smoke test", "chinese": "冒烟测试", "phonetic": "/sməʊk/", "example": "Perform smoke test.", "category": "测试用例"},
    {"id": "v133", "english": "integration test", "chinese": "集成测试", "phonetic": "/ˌɪntɪˈɡreɪʃən/", "example": "Run integration tests.", "category": "测试用例"},
    {"id": "v134", "english": "system test", "chinese": "系统测试", "phonetic": "/ˈsɪstəm/", "example": "Perform system test.", "category": "测试用例"},
    {"id": "v135", "english": "acceptance test", "chinese": "验收测试", "phonetic": "/əkˈseptəns/", "example": "Customer acceptance test.", "category": "测试用例"},
    {"id": "v136", "english": "boundary", "chinese": "边界", "phonetic": "/ˈbaʊndəri/", "example": "Test boundary conditions.", "category": "测试用例"},
    {"id": "v137", "english": "edge case", "chinese": "边缘情况", "phonetic": "/edʒ keɪs/", "example": "Consider edge cases.", "category": "测试用例"},
    {"id": "v138", "english": "coverage", "chinese": "覆盖率", "phonetic": "/ˈkʌvərɪdʒ/", "example": "Test coverage 85%.", "category": "测试用例"},
    {"id": "v139", "english": "automation", "chinese": "自动化", "phonetic": "/ˌɔːtəˈmeɪʃən/", "example": "Implement automation.", "category": "测试用例"},
    {"id": "v140", "english": "manual test", "chinese": "手动测试", "phonetic": "/ˈmænjuəl/", "example": "Manual test required.", "category": "测试用例"},
    {"id": "v141", "english": "HIL", "chinese": "硬件在环", "phonetic": "/eɪtʃ aɪ el/", "example": "HIL simulation test.", "category": "测试用例"},
    {"id": "v142", "english": "SIL", "chinese": "软件在环", "phonetic": "/es aɪ el/", "example": "SIL test environment.", "category": "测试用例"},
    {"id": "v143", "english": "bench", "chinese": "台架", "phonetic": "/bentʃ/", "example": "Test on bench first.", "category": "测试用例"},
    {"id": "v144", "english": "vehicle test", "chinese": "实车测试", "phonetic": "/ˈviːɪkəl/", "example": "Perform vehicle test.", "category": "测试用例"},
    {"id": "v145", "english": "road test", "chinese": "路试", "phonetic": "/rəʊd/", "example": "Road test required.", "category": "测试用例"},
    {"id": "v146", "english": "stress test", "chinese": "压力测试", "phonetic": "/stres/", "example": "Run stress test.", "category": "测试用例"},
    {"id": "v147", "english": "load test", "chinese": "负载测试", "phonetic": "/ləʊd/", "example": "Perform load test.", "category": "测试用例"},
    {"id": "v148", "english": "performance", "chinese": "性能", "phonetic": "/pəˈfɔːməns/", "example": "Check performance.", "category": "测试用例"},
    {"id": "v149", "english": "stability", "chinese": "稳定性", "phonetic": "/stəˈbɪləti/", "example": "Test stability.", "category": "测试用例"},
    {"id": "v150", "english": "durability", "chinese": "耐久性", "phonetic": "/ˌdjʊərəˈbɪləti/", "example": "Durability test.", "category": "测试用例"},
    {"id": "v151", "english": "scenario", "chinese": "场景", "phonetic": "/sɪˈnɑːriəʊ/", "example": "Test scenario.", "category": "测试用例"},
    {"id": "v152", "english": "use case", "chinese": "用例", "phonetic": "/juːs keɪs/", "example": "Define use case.", "category": "测试用例"},
    {"id": "v153", "english": "requirement", "chinese": "需求", "phonetic": "/rɪˈkwaɪəmənt/", "example": "Check requirement.", "category": "测试用例"},
    {"id": "v154", "english": "specification", "chinese": "规格", "phonetic": "/ˌspesɪfɪˈkeɪʃən/", "example": "Read specification.", "category": "测试用例"},
    {"id": "v155", "english": "traceability", "chinese": "可追溯性", "phonetic": "/ˌtreɪsəˈbɪləti/", "example": "Ensure traceability.", "category": "测试用例"},

    # Bug Reports (45)
    {"id": "v156", "english": "bug", "chinese": "缺陷", "phonetic": "/bʌɡ/", "example": "Found a bug.", "category": "缺陷报告"},
    {"id": "v157", "english": "defect", "chinese": "缺陷", "phonetic": "/ˈdiːfekt/", "example": "Report the defect.", "category": "缺陷报告"},
    {"id": "v158", "english": "issue", "chinese": "问题", "phonetic": "/ˈɪʃuː/", "example": "There is an issue.", "category": "缺陷报告"},
    {"id": "v159", "english": "severity", "chinese": "严重程度", "phonetic": "/sɪˈverəti/", "example": "Set severity critical.", "category": "缺陷报告"},
    {"id": "v160", "english": "priority", "chinese": "优先级", "phonetic": "/praɪˈɒrəti/", "example": "High priority bug.", "category": "缺陷报告"},
    {"id": "v161", "english": "critical", "chinese": "严重", "phonetic": "/ˈkrɪtɪkəl/", "example": "Critical bug found.", "category": "缺陷报告"},
    {"id": "v162", "english": "major", "chinese": "主要", "phonetic": "/ˈmeɪdʒə/", "example": "Major issue reported.", "category": "缺陷报告"},
    {"id": "v163", "english": "minor", "chinese": "次要", "phonetic": "/ˈmaɪnə/", "example": "Minor UI issue.", "category": "缺陷报告"},
    {"id": "v164", "english": "blocker", "chinese": "阻塞问题", "phonetic": "/ˈblɒkə/", "example": "This is a blocker.", "category": "缺陷报告"},
    {"id": "v165", "english": "reproduce", "chinese": "复现", "phonetic": "/ˌriːprəˈdjuːs/", "example": "Can you reproduce?", "category": "缺陷报告"},
    {"id": "v166", "english": "root cause", "chinese": "根本原因", "phonetic": "/ruːt kɔːz/", "example": "Analyze root cause.", "category": "缺陷报告"},
    {"id": "v167", "english": "workaround", "chinese": "临时方案", "phonetic": "/ˈwɜːkəraʊnd/", "example": "Is there workaround?", "category": "缺陷报告"},
    {"id": "v168", "english": "fix", "chinese": "修复", "phonetic": "/fɪks/", "example": "Fix the bug.", "category": "缺陷报告"},
    {"id": "v169", "english": "patch", "chinese": "补丁", "phonetic": "/pætʃ/", "example": "Apply the patch.", "category": "缺陷报告"},
    {"id": "v170", "english": "hotfix", "chinese": "热修复", "phonetic": "/ˈhɒtfɪks/", "example": "Release a hotfix.", "category": "缺陷报告"},
    {"id": "v171", "english": "reopen", "chinese": "重新打开", "phonetic": "/riːˈəʊpən/", "example": "Reopen the bug.", "category": "缺陷报告"},
    {"id": "v172", "english": "close", "chinese": "关闭", "phonetic": "/kləʊz/", "example": "Close the bug.", "category": "缺陷报告"},
    {"id": "v173", "english": "duplicate", "chinese": "重复", "phonetic": "/ˈdjuːplɪkət/", "example": "This is duplicate.", "category": "缺陷报告"},
    {"id": "v174", "english": "invalid", "chinese": "无效", "phonetic": "/ɪnˈvælɪd/", "example": "Mark as invalid.", "category": "缺陷报告"},
    {"id": "v175", "english": "log", "chinese": "日志", "phonetic": "/lɒɡ/", "example": "Attach log file.", "category": "缺陷报告"},
    {"id": "v176", "english": "screenshot", "chinese": "截图", "phonetic": "/ˈskriːnʃɒt/", "example": "Attach screenshot.", "category": "缺陷报告"},
    {"id": "v177", "english": "environment", "chinese": "环境", "phonetic": "/ɪnˈvaɪrənmənt/", "example": "Describe environment.", "category": "缺陷报告"},
    {"id": "v178", "english": "version", "chinese": "版本", "phonetic": "/ˈvɜːʃən/", "example": "Which version?", "category": "缺陷报告"},
    {"id": "v179", "english": "build", "chinese": "构建版本", "phonetic": "/bɪld/", "example": "Test on latest build.", "category": "缺陷报告"},
    {"id": "v180", "english": "release", "chinese": "发布", "phonetic": "/rɪˈliːs/", "example": "Release new version.", "category": "缺陷报告"},
    {"id": "v181", "english": "debug", "chinese": "调试", "phonetic": "/diːˈbʌɡ/", "example": "Debug the issue.", "category": "缺陷报告"},
    {"id": "v182", "english": "trace", "chinese": "追踪", "phonetic": "/treɪs/", "example": "Trace the problem.", "category": "缺陷报告"},
    {"id": "v183", "english": "analyze", "chinese": "分析", "phonetic": "/ˈænəlaɪz/", "example": "Analyze the log.", "category": "缺陷报告"},
    {"id": "v184", "english": "investigate", "chinese": "调查", "phonetic": "/ɪnˈvestɪɡeɪt/", "example": "Investigate issue.", "category": "缺陷报告"},
    {"id": "v185", "english": "solution", "chinese": "解决方案", "phonetic": "/səˈluːʃən/", "example": "Find a solution.", "category": "缺陷报告"},
    {"id": "v186", "english": "implement", "chinese": "实现", "phonetic": "/ˈɪmplɪment/", "example": "Implement the fix.", "category": "缺陷报告"},
    {"id": "v187", "english": "deploy", "chinese": "部署", "phonetic": "/dɪˈplɔɪ/", "example": "Deploy to production.", "category": "缺陷报告"},
    {"id": "v188", "english": "monitor", "chinese": "监控", "phonetic": "/ˈmɒnɪtə/", "example": "Monitor the update.", "category": "缺陷报告"},
    {"id": "v189", "english": "error code", "chinese": "错误码", "phonetic": "/ˈerə kəʊd/", "example": "Check error code.", "category": "缺陷报告"},
    {"id": "v190", "english": "exception", "chinese": "异常", "phonetic": "/ɪkˈsepʃən/", "example": "Handle exception.", "category": "缺陷报告"},
    {"id": "v191", "english": "warning", "chinese": "警告", "phonetic": "/ˈwɔːnɪŋ/", "example": "Warning message.", "category": "缺陷报告"},
    {"id": "v192", "english": "failure", "chinese": "失败", "phonetic": "/ˈfeɪljə/", "example": "Update failure.", "category": "缺陷报告"},
    {"id": "v193", "english": "success", "chinese": "成功", "phonetic": "/səkˈses/", "example": "Update success.", "category": "缺陷报告"},
    {"id": "v194", "english": "complete", "chinese": "完成", "phonetic": "/kəmˈpliːt/", "example": "Update complete.", "category": "缺陷报告"},
    {"id": "v195", "english": "abort", "chinese": "终止", "phonetic": "/əˈbɔːt/", "example": "Abort the update.", "category": "缺陷报告"},
    {"id": "v196", "english": "interrupt", "chinese": "中断", "phonetic": "/ˌɪntəˈrʌpt/", "example": "Update interrupted.", "category": "缺陷报告"},
    {"id": "v197", "english": "status", "chinese": "状态", "phonetic": "/ˈsteɪtəs/", "example": "Check update status.", "category": "缺陷报告"},
    {"id": "v198", "english": "feedback", "chinese": "反馈", "phonetic": "/ˈfiːdbæk/", "example": "Provide feedback.", "category": "缺陷报告"},
    {"id": "v199", "english": "confirm", "chinese": "确认", "phonetic": "/kənˈfɜːm/", "example": "Please confirm.", "category": "缺陷报告"},
    {"id": "v200", "english": "approve", "chinese": "批准", "phonetic": "/əˈpruːv/", "example": "Approve the update.", "category": "缺陷报告"},
]


# ==================== 扩展短语数据 (100+) ====================
OTA_PHRASES = [
    # Daily Standup (20)
    {"id": "p001", "english": "Good morning everyone, let's start the standup.", "chinese": "大家早上好，我们开始站会吧。", "scenario": "日常站会"},
    {"id": "p002", "english": "Yesterday I worked on OTA testing.", "chinese": "昨天我做了OTA测试工作。", "scenario": "日常站会"},
    {"id": "p003", "english": "Today I will continue the regression test.", "chinese": "今天我会继续回归测试。", "scenario": "日常站会"},
    {"id": "p004", "english": "I'm blocked by the server issue.", "chinese": "我被服务器问题阻塞了。", "scenario": "日常站会"},
    {"id": "p005", "english": "No blockers from my side.", "chinese": "我这边没有阻塞问题。", "scenario": "日常站会"},
    {"id": "p006", "english": "I need help with the test environment.", "chinese": "我需要测试环境方面的帮助。", "scenario": "日常站会"},
    {"id": "p007", "english": "The test progress is on track.", "chinese": "测试进度正常。", "scenario": "日常站会"},
    {"id": "p008", "english": "We found 3 bugs yesterday.", "chinese": "昨天我们发现了3个bug。", "scenario": "日常站会"},
    {"id": "p009", "english": "All test cases passed.", "chinese": "所有测试用例都通过了。", "scenario": "日常站会"},
    {"id": "p010", "english": "Can we have a quick sync?", "chinese": "我们可以快速同步一下吗？", "scenario": "日常站会"},
    {"id": "p011", "english": "Let me share my screen.", "chinese": "让我分享一下屏幕。", "scenario": "日常站会"},
    {"id": "p012", "english": "Can you hear me clearly?", "chinese": "你能听清楚吗？", "scenario": "日常站会"},
    {"id": "p013", "english": "Sorry, I was on mute.", "chinese": "抱歉，我刚才静音了。", "scenario": "日常站会"},
    {"id": "p014", "english": "Let's take this offline.", "chinese": "我们线下讨论这个问题。", "scenario": "日常站会"},
    {"id": "p015", "english": "I will send the meeting notes.", "chinese": "我会发送会议纪要。", "scenario": "日常站会"},
    {"id": "p016", "english": "Any other topics to discuss?", "chinese": "还有其他话题要讨论吗？", "scenario": "日常站会"},
    {"id": "p017", "english": "Thanks everyone, meeting adjourned.", "chinese": "谢谢大家，会议结束。", "scenario": "日常站会"},
    {"id": "p018", "english": "I will update the test report today.", "chinese": "今天我会更新测试报告。", "scenario": "日常站会"},
    {"id": "p019", "english": "The deadline is approaching.", "chinese": "截止日期快到了。", "scenario": "日常站会"},
    {"id": "p020", "english": "We need to speed up the testing.", "chinese": "我们需要加快测试进度。", "scenario": "日常站会"},
    
    # Bug Report (25)
    {"id": "p021", "english": "I found a bug in the OTA download process.", "chinese": "我在OTA下载过程中发现了一个bug。", "scenario": "缺陷报告"},
    {"id": "p022", "english": "The OTA update failed at 80% progress.", "chinese": "OTA升级在80%进度时失败了。", "scenario": "缺陷报告"},
    {"id": "p023", "english": "Can you reproduce this issue?", "chinese": "你能复现这个问题吗？", "scenario": "缺陷报告"},
    {"id": "p024", "english": "I can reproduce it every time.", "chinese": "我每次都能复现。", "scenario": "缺陷报告"},
    {"id": "p025", "english": "It's an intermittent issue.", "chinese": "这是一个偶发问题。", "scenario": "缺陷报告"},
    {"id": "p026", "english": "Please attach the log file.", "chinese": "请附上日志文件。", "scenario": "缺陷报告"},
    {"id": "p027", "english": "What's the software version?", "chinese": "软件版本是什么？", "scenario": "缺陷报告"},
    {"id": "p028", "english": "The root cause is identified.", "chinese": "根本原因已确定。", "scenario": "缺陷报告"},
    {"id": "p029", "english": "The fix will be in the next release.", "chinese": "修复会在下个版本发布。", "scenario": "缺陷报告"},
    {"id": "p030", "english": "This is a critical bug.", "chinese": "这是一个严重bug。", "scenario": "缺陷报告"},
    {"id": "p031", "english": "Please verify the fix.", "chinese": "请验证修复。", "scenario": "缺陷报告"},
    {"id": "p032", "english": "The bug is confirmed fixed.", "chinese": "bug已确认修复。", "scenario": "缺陷报告"},
    {"id": "p033", "english": "I will close this bug.", "chinese": "我会关闭这个bug。", "scenario": "缺陷报告"},
    {"id": "p034", "english": "We need a hotfix for this.", "chinese": "我们需要一个热修复。", "scenario": "缺陷报告"},
    {"id": "p035", "english": "Is there a workaround?", "chinese": "有临时解决方案吗？", "scenario": "缺陷报告"},
    {"id": "p036", "english": "The issue occurs randomly.", "chinese": "这个问题随机出现。", "scenario": "缺陷报告"},
    {"id": "p037", "english": "I cannot reproduce it anymore.", "chinese": "我无法再复现了。", "scenario": "缺陷报告"},
    {"id": "p038", "english": "The issue is environment specific.", "chinese": "这个问题是环境相关的。", "scenario": "缺陷报告"},
    {"id": "p039", "english": "Please provide more details.", "chinese": "请提供更多细节。", "scenario": "缺陷报告"},
    {"id": "p040", "english": "What are the steps to reproduce?", "chinese": "复现步骤是什么？", "scenario": "缺陷报告"},
    {"id": "p041", "english": "The bug is assigned to developer.", "chinese": "bug已分配给开发人员。", "scenario": "缺陷报告"},
    {"id": "p042", "english": "This is a known issue.", "chinese": "这是一个已知问题。", "scenario": "缺陷报告"},
    {"id": "p043", "english": "The bug has been reopened.", "chinese": "bug已被重新打开。", "scenario": "缺陷报告"},
    {"id": "p044", "english": "Please update the bug status.", "chinese": "请更新bug状态。", "scenario": "缺陷报告"},
    {"id": "p045", "english": "The severity should be changed to critical.", "chinese": "严重程度应该改为严重。", "scenario": "缺陷报告"},
    
    # Test Planning (20)
    {"id": "p046", "english": "Let's review the test plan.", "chinese": "我们来评审测试计划。", "scenario": "测试计划"},
    {"id": "p047", "english": "What's the test scope?", "chinese": "测试范围是什么？", "scenario": "测试计划"},
    {"id": "p048", "english": "We need to add more test cases.", "chinese": "我们需要添加更多测试用例。", "scenario": "测试计划"},
    {"id": "p049", "english": "The test coverage is not enough.", "chinese": "测试覆盖率不够。", "scenario": "测试计划"},
    {"id": "p050", "english": "When is the test deadline?", "chinese": "测试截止日期是什么时候？", "scenario": "测试计划"},
    {"id": "p051", "english": "We need more test resources.", "chinese": "我们需要更多测试资源。", "scenario": "测试计划"},
    {"id": "p052", "english": "The test environment is ready.", "chinese": "测试环境已准备好。", "scenario": "测试计划"},
    {"id": "p053", "english": "We should do smoke test first.", "chinese": "我们应该先做冒烟测试。", "scenario": "测试计划"},
    {"id": "p054", "english": "How many test cycles do we need?", "chinese": "我们需要多少轮测试？", "scenario": "测试计划"},
    {"id": "p055", "english": "The test report is ready.", "chinese": "测试报告已准备好。", "scenario": "测试计划"},
    {"id": "p056", "english": "Let's define the acceptance criteria.", "chinese": "我们来定义验收标准。", "scenario": "测试计划"},
    {"id": "p057", "english": "We need to test all ECUs.", "chinese": "我们需要测试所有ECU。", "scenario": "测试计划"},
    {"id": "p058", "english": "The test data is prepared.", "chinese": "测试数据已准备好。", "scenario": "测试计划"},
    {"id": "p059", "english": "We should automate this test.", "chinese": "我们应该自动化这个测试。", "scenario": "测试计划"},
    {"id": "p060", "english": "Manual testing is required.", "chinese": "需要手动测试。", "scenario": "测试计划"},
    {"id": "p061", "english": "Let's estimate the test effort.", "chinese": "我们来估算测试工作量。", "scenario": "测试计划"},
    {"id": "p062", "english": "We passed all critical tests.", "chinese": "我们通过了所有关键测试。", "scenario": "测试计划"},
    {"id": "p063", "english": "Some tests are still pending.", "chinese": "一些测试仍在进行中。", "scenario": "测试计划"},
    {"id": "p064", "english": "The release is approved.", "chinese": "发布已批准。", "scenario": "测试计划"},
    {"id": "p065", "english": "We need to prioritize the test cases.", "chinese": "我们需要确定测试用例优先级。", "scenario": "测试计划"},

    # Technical Discussion (20)
    {"id": "p066", "english": "The download speed is too slow.", "chinese": "下载速度太慢了。", "scenario": "技术讨论"},
    {"id": "p067", "english": "We need to optimize the algorithm.", "chinese": "我们需要优化算法。", "scenario": "技术讨论"},
    {"id": "p068", "english": "What protocol are we using?", "chinese": "我们使用什么协议？", "scenario": "技术讨论"},
    {"id": "p069", "english": "The checksum verification failed.", "chinese": "校验和验证失败了。", "scenario": "技术讨论"},
    {"id": "p070", "english": "We should use delta update.", "chinese": "我们应该使用差分升级。", "scenario": "技术讨论"},
    {"id": "p071", "english": "The ECU response is incorrect.", "chinese": "ECU响应不正确。", "scenario": "技术讨论"},
    {"id": "p072", "english": "Let me check the diagnostic log.", "chinese": "让我检查一下诊断日志。", "scenario": "技术讨论"},
    {"id": "p073", "english": "The security access failed.", "chinese": "安全访问失败了。", "scenario": "技术讨论"},
    {"id": "p074", "english": "The network connection is unstable.", "chinese": "网络连接不稳定。", "scenario": "技术讨论"},
    {"id": "p075", "english": "Let's analyze the failure case.", "chinese": "我们来分析失败案例。", "scenario": "技术讨论"},
    {"id": "p076", "english": "The update package is corrupted.", "chinese": "升级包已损坏。", "scenario": "技术讨论"},
    {"id": "p077", "english": "We need to verify the signature.", "chinese": "我们需要验证签名。", "scenario": "技术讨论"},
    {"id": "p078", "english": "The certificate has expired.", "chinese": "证书已过期。", "scenario": "技术讨论"},
    {"id": "p079", "english": "Let's check the server status.", "chinese": "我们检查一下服务器状态。", "scenario": "技术讨论"},
    {"id": "p080", "english": "The API response is incorrect.", "chinese": "API响应不正确。", "scenario": "技术讨论"},
    {"id": "p081", "english": "We need to handle edge cases.", "chinese": "我们需要处理边缘情况。", "scenario": "技术讨论"},
    {"id": "p082", "english": "The battery level is too low.", "chinese": "电池电量太低。", "scenario": "技术讨论"},
    {"id": "p083", "english": "The vehicle is not in park mode.", "chinese": "车辆不在驻车模式。", "scenario": "技术讨论"},
    {"id": "p084", "english": "We should add retry mechanism.", "chinese": "我们应该添加重试机制。", "scenario": "技术讨论"},
    {"id": "p085", "english": "The download was interrupted.", "chinese": "下载被中断了。", "scenario": "技术讨论"},
    
    # Email Communication (20)
    {"id": "p086", "english": "Dear team, please find the test report attached.", "chinese": "团队好，请查收附件中的测试报告。", "scenario": "邮件沟通"},
    {"id": "p087", "english": "Could you please review the test cases?", "chinese": "请您评审一下测试用例好吗？", "scenario": "邮件沟通"},
    {"id": "p088", "english": "I have updated the bug status.", "chinese": "我已更新bug状态。", "scenario": "邮件沟通"},
    {"id": "p089", "english": "Please let me know if you have any questions.", "chinese": "如有任何问题请告诉我。", "scenario": "邮件沟通"},
    {"id": "p090", "english": "Thanks for your quick response.", "chinese": "感谢您的快速回复。", "scenario": "邮件沟通"},
    {"id": "p091", "english": "I will follow up on this issue.", "chinese": "我会跟进这个问题。", "scenario": "邮件沟通"},
    {"id": "p092", "english": "The test is completed successfully.", "chinese": "测试已成功完成。", "scenario": "邮件沟通"},
    {"id": "p093", "english": "We need your approval to proceed.", "chinese": "我们需要您的批准才能继续。", "scenario": "邮件沟通"},
    {"id": "p094", "english": "Best regards.", "chinese": "此致敬礼。", "scenario": "邮件沟通"},
    {"id": "p095", "english": "Looking forward to your feedback.", "chinese": "期待您的反馈。", "scenario": "邮件沟通"},
    {"id": "p096", "english": "Please find the release notes below.", "chinese": "请查看以下发布说明。", "scenario": "邮件沟通"},
    {"id": "p097", "english": "I apologize for the delay.", "chinese": "对于延迟我深表歉意。", "scenario": "邮件沟通"},
    {"id": "p098", "english": "Please escalate this issue.", "chinese": "请升级这个问题。", "scenario": "邮件沟通"},
    {"id": "p099", "english": "I will send the update by EOD.", "chinese": "我会在今天结束前发送更新。", "scenario": "邮件沟通"},
    {"id": "p100", "english": "Please confirm the test schedule.", "chinese": "请确认测试计划。", "scenario": "邮件沟通"},
    {"id": "p101", "english": "The meeting has been rescheduled.", "chinese": "会议已重新安排。", "scenario": "邮件沟通"},
    {"id": "p102", "english": "Please review and provide your comments.", "chinese": "请评审并提供您的意见。", "scenario": "邮件沟通"},
    {"id": "p103", "english": "I have attached the updated document.", "chinese": "我已附上更新的文档。", "scenario": "邮件沟通"},
    {"id": "p104", "english": "Please acknowledge receipt of this email.", "chinese": "请确认收到此邮件。", "scenario": "邮件沟通"},
    {"id": "p105", "english": "Thank you for your cooperation.", "chinese": "感谢您的配合。", "scenario": "邮件沟通"},
]

# ==================== 自定义词汇存储 ====================
# 使用用户目录存储数据，避免权限问题
import os
DATA_DIR = Path(os.path.expanduser("~")) / ".ota_english"
DATA_DIR.mkdir(exist_ok=True)
CUSTOM_DATA_FILE = DATA_DIR / "custom_vocabulary.json"
PROGRESS_FILE = DATA_DIR / "progress.json"

def load_custom_data():
    """加载自定义词汇"""
    try:
        if CUSTOM_DATA_FILE.exists():
            with open(CUSTOM_DATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        st.warning(f"加载自定义词汇失败: {e}")
    return {"vocabulary": [], "phrases": []}

def save_custom_data(data):
    """保存自定义词汇"""
    try:
        with open(CUSTOM_DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        st.error(f"保存失败: {e}")
        return False

# ==================== 文档解析功能 ====================
def extract_words_from_text(text):
    """从文本中提取英文单词（去重，保留原始大小写）"""
    # 提取所有英文单词（包括带下划线和数字的技术术语）
    words = re.findall(r'\b[a-zA-Z][a-zA-Z0-9_\-]*[a-zA-Z0-9]\b', text)
    # 去重，保留第一次出现的大小写形式
    seen = set()
    unique_words = []
    for w in words:
        lower = w.lower()
        if lower not in seen and len(w) > 2:
            seen.add(lower)
            unique_words.append(w)
    # 过滤常见词（扩展列表）
    common_words = {
        # 冠词、代词、介词、连词
        'the', 'and', 'for', 'are', 'but', 'not', 'you', 'all', 'can', 'had', 'her', 'was', 'one', 'our', 'out', 
        'has', 'have', 'been', 'will', 'more', 'when', 'who', 'way', 'may', 'its', 'than', 'them', 'then', 'into', 
        'some', 'could', 'other', 'which', 'their', 'there', 'would', 'about', 'these', 'from', 'with', 'this', 
        'that', 'what', 'were', 'they', 'each', 'make', 'like', 'just', 'over', 'such', 'also', 'back', 'after', 
        'most', 'only', 'come', 'made', 'find', 'here', 'many', 'where', 'does', 'being', 'under', 'last', 'right', 
        'still', 'must', 'own', 'through', 'before', 'same', 'should', 'well', 'between', 'each', 'because', 'very', 
        'without', 'again', 'off', 'might', 'while', 'never', 'below', 'next', 'few', 'those', 'always', 'both', 
        'how', 'why', 'any', 'new', 'work', 'first', 'second', 'third', 'part', 'take', 'get', 'place', 'number', 
        'year', 'day', 'good', 'give', 'use', 'say', 'see', 'know', 'want', 'look', 'think', 'time', 'now', 'people', 
        'even', 'thing', 'man', 'world', 'life', 'hand', 'high', 'old', 'great', 'big', 'small', 'large', 'long', 
        'little', 'left', 'early', 'young', 'important', 'public', 'bad', 'able', 'shall', 'need',
        # 常见动词
        'show', 'display', 'click', 'select', 'enter', 'input', 'output', 'create', 'delete', 'update', 'read', 
        'write', 'save', 'load', 'send', 'receive', 'connect', 'disconnect', 'start', 'stop', 'open', 'close',
        'enable', 'disable', 'true', 'false', 'null', 'none', 'default', 'custom',
        # 常见名词（保留技术相关的）
        'user', 'system', 'data', 'information', 'function', 'feature', 'page', 'button', 'file', 'name', 'type', 
        'value', 'list', 'item', 'table', 'row', 'column', 'field', 'form', 'text', 'image', 'icon', 'menu', 
        'option', 'setting', 'config', 'parameter', 'result', 'status', 'state', 'mode', 'level', 'size', 
        'width', 'height', 'color', 'style', 'format', 'content', 'title', 'description', 'note', 'comment', 
        'message', 'error', 'warning', 'success', 'fail', 'request', 'response'
    }
    words = [w for w in unique_words if w.lower() not in common_words]
    return words

def extract_sentences_from_text(text):
    """从文本中提取英文句子（去重）"""
    # 按句号、问号、感叹号分割
    sentences = re.split(r'[.!?]+', text)
    # 清理并过滤
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20 and len(s.strip()) < 300]
    # 只保留主要是英文的句子，并去重
    seen = set()
    english_sentences = []
    for s in sentences:
        english_chars = len(re.findall(r'[a-zA-Z]', s))
        lower_s = s.lower()
        if english_chars > len(s) * 0.5 and lower_s not in seen:  # 超过50%是英文字符
            seen.add(lower_s)
            english_sentences.append(s)
    return english_sentences  # 返回所有句子，不限制数量

# ==================== 数据存储 ====================
# 使用用户目录中的 PROGRESS_FILE（已在上面定义）

def load_progress():
    try:
        if PROGRESS_FILE.exists():
            with open(PROGRESS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
    except:
        pass
    return {"mastery": {}, "favorites": [], "quiz_history": [], "streak": 0, "last_study": None}

def save_progress(progress):
    try:
        with open(PROGRESS_FILE, 'w', encoding='utf-8') as f:
            json.dump(progress, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"保存进度失败: {e}")

# ==================== 获取所有词汇（内置+自定义）====================
def get_all_vocabulary():
    custom = load_custom_data()
    return OTA_VOCABULARY + custom.get("vocabulary", [])

def get_all_phrases():
    custom = load_custom_data()
    return OTA_PHRASES + custom.get("phrases", [])

# ==================== Session State ====================
if 'progress' not in st.session_state:
    st.session_state.progress = load_progress()
if 'current_page' not in st.session_state:
    st.session_state.current_page = "首页"
if 'flashcard_index' not in st.session_state:
    st.session_state.flashcard_index = 0
if 'flashcard_flipped' not in st.session_state:
    st.session_state.flashcard_flipped = False
if 'quiz_questions' not in st.session_state:
    st.session_state.quiz_questions = []
if 'quiz_index' not in st.session_state:
    st.session_state.quiz_index = 0
if 'quiz_score' not in st.session_state:
    st.session_state.quiz_score = 0
if 'quiz_answered' not in st.session_state:
    st.session_state.quiz_answered = False


# ==================== 语音按钮辅助函数 ====================
def create_speak_button(english, chinese, button_id, button_text="🔊 朗读"):
    """创建可点击的语音按钮"""
    # 转义特殊字符
    english_safe = english.replace("'", "\\'").replace('"', '\\"').replace('\n', ' ')
    chinese_safe = chinese.replace("'", "\\'").replace('"', '\\"').replace('\n', ' ')
    
    html_code = f"""
    <button 
        onclick="speakWord_{button_id}()" 
        style="background: #4CAF50; color: white; border: none; padding: 8px 16px; 
               border-radius: 8px; cursor: pointer; font-size: 14px; margin: 5px 2px;
               transition: all 0.3s;"
        onmouseover="this.style.background='#45a049'; this.style.transform='scale(1.05)';"
        onmouseout="this.style.background='#4CAF50'; this.style.transform='scale(1)';">
        {button_text}
    </button>
    <script>
    function speakWord_{button_id}() {{
        if ('speechSynthesis' in window) {{
            window.speechSynthesis.cancel();
            
            const utterance1 = new SpeechSynthesisUtterance('{english_safe}');
            utterance1.lang = 'en-US';
            utterance1.rate = 0.7;  // 降低速度，更清晰
            utterance1.pitch = 1.0;  // 标准音调
            utterance1.volume = 1.0;  // 最大音量
            
            const utterance2 = new SpeechSynthesisUtterance('{chinese_safe}');
            utterance2.lang = 'zh-CN';
            utterance2.rate = 0.75;  // 中文稍快一点
            utterance2.pitch = 1.0;
            utterance2.volume = 1.0;
            
            window.speechSynthesis.speak(utterance1);
            
            utterance1.onend = function() {{
                setTimeout(function() {{
                    window.speechSynthesis.speak(utterance2);
                }}, 500);  // 增加停顿时间
            }};
        }} else {{
            alert('您的浏览器不支持语音功能');
        }}
    }}
    </script>
    """
    components.html(html_code, height=50)

def create_speak_button_english_only(text, button_id, button_text="🔊 朗读例句"):
    """创建只读英文的语音按钮"""
    text_safe = text.replace("'", "\\'").replace('"', '\\"').replace('\n', ' ')
    
    html_code = f"""
    <button 
        onclick="speakEnglish_{button_id}()" 
        style="background: #4CAF50; color: white; border: none; padding: 8px 16px; 
               border-radius: 8px; cursor: pointer; font-size: 14px; margin: 5px 2px;
               transition: all 0.3s;"
        onmouseover="this.style.background='#45a049'; this.style.transform='scale(1.05)';"
        onmouseout="this.style.background='#4CAF50'; this.style.transform='scale(1)';">
        {button_text}
    </button>
    <script>
    function speakEnglish_{button_id}() {{
        if ('speechSynthesis' in window) {{
            window.speechSynthesis.cancel();
            
            const utterance = new SpeechSynthesisUtterance('{text_safe}');
            utterance.lang = 'en-US';
            utterance.rate = 0.7;  // 降低速度，更清晰
            utterance.pitch = 1.0;  // 标准音调
            utterance.volume = 1.0;  // 最大音量
            
            window.speechSynthesis.speak(utterance);
        }} else {{
            alert('您的浏览器不支持语音功能');
        }}
    }}
    </script>
    """
    components.html(html_code, height=50)

# ==================== 页面函数 ====================

def show_home():
    """首页"""
    st.markdown("## 📚 OTA英语学习 V2.0")
    st.markdown("---")
    
    vocab = get_all_vocabulary()
    phrases = get_all_phrases()
    progress = st.session_state.progress
    mastery = progress.get("mastery", {})
    
    total_words = len(vocab)
    mastered = sum(1 for v in mastery.values() if v >= 80)
    learning = sum(1 for v in mastery.values() if 0 < v < 80)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📖 总词汇", f"{total_words}")
    with col2:
        st.metric("✅ 已掌握", f"{mastered}")
    with col3:
        st.metric("💬 短语", f"{len(phrases)}")
    
    st.markdown("---")
    st.markdown("### 🚀 快速开始")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📖 学习词汇", use_container_width=True):
            st.session_state.current_page = "词汇"
            st.rerun()
    with col2:
        if st.button("💬 学习短语", use_container_width=True):
            st.session_state.current_page = "短语"
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("🎴 闪卡练习", use_container_width=True):
            st.session_state.current_page = "闪卡"
            st.rerun()
    with col2:
        if st.button("📝 开始测验", use_container_width=True):
            st.session_state.current_page = "测验"
            st.rerun()
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 导入文档", use_container_width=True):
            st.session_state.current_page = "导入"
            st.rerun()
    with col2:
        if st.button("⚙️ 管理词库", use_container_width=True):
            st.session_state.current_page = "管理"
            st.rerun()
    
    st.markdown("---")
    st.markdown("### 💡 今日推荐")
    word = random.choice(vocab)
    
    st.markdown(f"""
    <div class="word-card">
        <div style="font-size: 24px; font-weight: bold;">{word['english']}</div>
        <div style="font-size: 14px; opacity: 0.9;">{word['phonetic']}</div>
        <div style="font-size: 20px; margin: 10px 0;">{word['chinese']}</div>
        <div style="font-size: 13px; opacity: 0.8;">💡 {word['example']}</div>
    </div>
    """, unsafe_allow_html=True)
    
    # 添加语音按钮
    create_speak_button(word['english'], word['chinese'], "home_recommend")


def show_vocabulary():
    """词汇学习"""
    st.markdown("## 📖 词汇学习")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.rerun()
    
    st.markdown("---")
    vocab = get_all_vocabulary()
    
    categories = ["全部"] + list(set(w["category"] for w in vocab))
    selected_category = st.selectbox("选择分类", categories)
    search = st.text_input("🔍 搜索", placeholder="输入英文或中文...")
    
    words = vocab
    if selected_category != "全部":
        words = [w for w in words if w["category"] == selected_category]
    if search:
        search_lower = search.lower()
        words = [w for w in words if search_lower in w["english"].lower() or search in w["chinese"]]
    
    # 翻页设置
    ITEMS_PER_PAGE = 30
    total_words = len(words)
    total_pages = max(1, (total_words + ITEMS_PER_PAGE - 1) // ITEMS_PER_PAGE)
    
    # 初始化页码
    if 'vocab_page' not in st.session_state:
        st.session_state.vocab_page = 1
    
    # 确保页码有效
    if st.session_state.vocab_page > total_pages:
        st.session_state.vocab_page = 1
    
    current_page = st.session_state.vocab_page
    start_idx = (current_page - 1) * ITEMS_PER_PAGE
    end_idx = min(start_idx + ITEMS_PER_PAGE, total_words)
    
    st.markdown(f"共 **{total_words}** 个词汇 | 第 **{current_page}** / **{total_pages}** 页 | 显示 {start_idx + 1}-{end_idx}")
    
    # 翻页控制
    col1, col2, col3, col4, col5 = st.columns([1, 1, 2, 1, 1])
    with col1:
        if st.button("⏮️ 首页", disabled=current_page == 1):
            st.session_state.vocab_page = 1
            st.rerun()
    with col2:
        if st.button("◀️ 上一页", disabled=current_page == 1):
            st.session_state.vocab_page = current_page - 1
            st.rerun()
    with col3:
        # 快速跳转
        page_input = st.number_input("跳转到", min_value=1, max_value=total_pages, value=current_page, label_visibility="collapsed")
        if page_input != current_page:
            st.session_state.vocab_page = page_input
            st.rerun()
    with col4:
        if st.button("下一页 ▶️", disabled=current_page == total_pages):
            st.session_state.vocab_page = current_page + 1
            st.rerun()
    with col5:
        if st.button("末页 ⏭️", disabled=current_page == total_pages):
            st.session_state.vocab_page = total_pages
            st.rerun()
    
    st.markdown("---")
    
    # 显示当前页的单词
    for idx, word in enumerate(words[start_idx:end_idx]):
        with st.expander(f"**{word['english']}** - {word['chinese']}"):
            # 添加语音按钮
            button_id = f"vocab_{start_idx + idx}"
            create_speak_button(word['english'], word['chinese'], button_id, "🔊 朗读单词")
            
            if word.get('phonetic'):
                st.markdown(f"**发音:** {word['phonetic']}")
            st.markdown(f"**分类:** {word['category']}")
            if word.get('example'):
                st.markdown(f"**例句:** {word['example']}")
                # 例句朗读按钮
                example_button_id = f"example_{start_idx + idx}"
                create_speak_button_english_only(word['example'], example_button_id, "🔊 朗读例句")
            mastery = st.session_state.progress.get("mastery", {}).get(word["id"], 0)
            st.progress(mastery / 100)
            st.caption(f"掌握程度: {mastery}%")
    
    # 底部翻页
    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("◀️ 上页", key="prev_bottom", disabled=current_page == 1):
            st.session_state.vocab_page = current_page - 1
            st.rerun()
    with col2:
        st.markdown(f"<center>第 {current_page} / {total_pages} 页</center>", unsafe_allow_html=True)
    with col3:
        if st.button("下页 ▶️", key="next_bottom", disabled=current_page == total_pages):
            st.session_state.vocab_page = current_page + 1
            st.rerun()


def show_phrases():
    """短语学习"""
    st.markdown("## 💬 短语学习")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.rerun()
    
    st.markdown("---")
    phrases = get_all_phrases()
    
    scenarios = ["全部"] + list(set(p["scenario"] for p in phrases))
    selected_scenario = st.selectbox("选择场景", scenarios)
    
    items = phrases
    if selected_scenario != "全部":
        items = [p for p in items if p["scenario"] == selected_scenario]
    
    # 翻页设置
    ITEMS_PER_PAGE = 20
    total_items = len(items)
    total_pages = max(1, (total_items + ITEMS_PER_PAGE - 1) // ITEMS_PER_PAGE)
    
    if 'phrase_page' not in st.session_state:
        st.session_state.phrase_page = 1
    if st.session_state.phrase_page > total_pages:
        st.session_state.phrase_page = 1
    
    current_page = st.session_state.phrase_page
    start_idx = (current_page - 1) * ITEMS_PER_PAGE
    end_idx = min(start_idx + ITEMS_PER_PAGE, total_items)
    
    st.markdown(f"共 **{total_items}** 个短语 | 第 **{current_page}** / **{total_pages}** 页")
    
    # 翻页控制
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("◀️ 上一页", key="phrase_prev", disabled=current_page == 1):
            st.session_state.phrase_page = current_page - 1
            st.rerun()
    with col2:
        page_input = st.number_input("跳转", min_value=1, max_value=total_pages, value=current_page, key="phrase_jump", label_visibility="collapsed")
        if page_input != current_page:
            st.session_state.phrase_page = page_input
            st.rerun()
    with col3:
        if st.button("下一页 ▶️", key="phrase_next", disabled=current_page == total_pages):
            st.session_state.phrase_page = current_page + 1
            st.rerun()
    
    st.markdown("---")
    
    for idx, phrase in enumerate(items[start_idx:end_idx]):
        st.markdown(f"""
        <div class="phrase-card">
            <div style="font-size: 15px; margin-bottom: 8px;">{phrase['english']}</div>
            <div style="font-size: 14px; opacity: 0.9;">{phrase['chinese']}</div>
            <div style="font-size: 12px; opacity: 0.7; margin-top: 5px;">📍 {phrase['scenario']}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # 添加语音按钮
        button_id = f"phrase_{start_idx + idx}"
        create_speak_button(phrase['english'], phrase['chinese'], button_id)


def show_flashcards():
    """闪卡练习"""
    st.markdown("## 🎴 闪卡练习")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.rerun()
    
    st.markdown("---")
    vocab = get_all_vocabulary()
    
    if 'flashcard_words' not in st.session_state or st.button("🔄 重新开始"):
        st.session_state.flashcard_words = random.sample(vocab, min(50, len(vocab)))
        st.session_state.flashcard_index = 0
        st.session_state.flashcard_flipped = False
        st.rerun()
    
    words = st.session_state.flashcard_words
    index = st.session_state.flashcard_index
    
    if index >= len(words):
        st.success("🎉 恭喜！你已完成所有闪卡！")
        if st.button("重新开始"):
            st.session_state.flashcard_index = 0
            st.rerun()
        return
    
    word = words[index]
    st.progress((index + 1) / len(words))
    st.caption(f"进度: {index + 1} / {len(words)}")
    
    if st.session_state.flashcard_flipped:
        st.markdown(f"""
        <div class="flashcard">
            <div style="font-size: 22px; margin-bottom: 10px;">{word['english']}</div>
            <div style="font-size: 14px; opacity: 0.9;">{word['phonetic']}</div>
            <div style="font-size: 24px; margin: 15px 0;">{word['chinese']}</div>
            <div style="font-size: 13px; opacity: 0.8;">💡 {word['example']}</div>
        </div>
        """, unsafe_allow_html=True)
        
        # 闪卡翻转后自动播放
        english_safe = word['english'].replace("'", "\\'")
        chinese_safe = word['chinese'].replace("'", "\\'")
        
        flashcard_html = f"""
        <script>
        (function() {{
            if ('speechSynthesis' in window) {{
                window.speechSynthesis.cancel();
                
                const utterance1 = new SpeechSynthesisUtterance('{english_safe}');
                utterance1.lang = 'en-US';
                utterance1.rate = 0.7;  // 降低速度，更清晰
                utterance1.pitch = 1.0;  // 标准音调
                utterance1.volume = 1.0;  // 最大音量
                
                const utterance2 = new SpeechSynthesisUtterance('{chinese_safe}');
                utterance2.lang = 'zh-CN';
                utterance2.rate = 0.75;  // 中文稍快一点
                utterance2.pitch = 1.0;
                utterance2.volume = 1.0;
                
                window.speechSynthesis.speak(utterance1);
                utterance1.onend = function() {{
                    setTimeout(function() {{
                        window.speechSynthesis.speak(utterance2);
                    }}, 500);  // 增加停顿时间
                }};
            }}
        }})();
        </script>
        """
        components.html(flashcard_html, height=0)
    else:
        st.markdown(f"""
        <div class="flashcard">
            <div style="font-size: 28px; font-weight: bold;">{word['english']}</div>
            <div style="font-size: 14px; margin-top: 8px; opacity: 0.8;">{word['phonetic']}</div>
            <div style="margin-top: 25px; font-size: 13px;">👆 点击翻转查看答案</div>
        </div>
        """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if not st.session_state.flashcard_flipped:
            if st.button("🔄 翻转", use_container_width=True):
                st.session_state.flashcard_flipped = True
                st.rerun()
    with col2:
        if st.session_state.flashcard_flipped:
            if st.button("❌ 不认识", use_container_width=True):
                mastery = st.session_state.progress.get("mastery", {})
                current = mastery.get(word["id"], 0)
                mastery[word["id"]] = max(0, current - 10)
                st.session_state.progress["mastery"] = mastery
                save_progress(st.session_state.progress)
                st.session_state.flashcard_index += 1
                st.session_state.flashcard_flipped = False
                st.rerun()
    with col3:
        if st.session_state.flashcard_flipped:
            if st.button("✅ 认识", use_container_width=True):
                mastery = st.session_state.progress.get("mastery", {})
                current = mastery.get(word["id"], 0)
                mastery[word["id"]] = min(100, current + 15)
                st.session_state.progress["mastery"] = mastery
                save_progress(st.session_state.progress)
                st.session_state.flashcard_index += 1
                st.session_state.flashcard_flipped = False
                st.rerun()


def show_quiz():
    """测验"""
    st.markdown("## 📝 词汇测验")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.session_state.quiz_questions = []
        st.rerun()
    
    st.markdown("---")
    vocab = get_all_vocabulary()
    
    if not st.session_state.quiz_questions:
        categories = ["全部"] + list(set(w["category"] for w in vocab))
        selected_category = st.selectbox("选择分类", categories)
        num_questions = st.slider("题目数量", 5, 20, 10)
        
        if st.button("🚀 开始测验", use_container_width=True):
            words = vocab
            if selected_category != "全部":
                words = [w for w in words if w["category"] == selected_category]
            
            if len(words) < 4:
                st.error("词汇太少，请选择其他分类")
                return
            
            selected = random.sample(words, min(num_questions, len(words)))
            questions = []
            for word in selected:
                wrong = random.sample([w for w in words if w["id"] != word["id"]], min(3, len(words)-1))
                options = [word["chinese"]] + [w["chinese"] for w in wrong]
                random.shuffle(options)
                questions.append({"word": word, "options": options, "correct": word["chinese"]})
            
            st.session_state.quiz_questions = questions
            st.session_state.quiz_index = 0
            st.session_state.quiz_score = 0
            st.session_state.quiz_answered = False
            st.rerun()
        return
    
    questions = st.session_state.quiz_questions
    index = st.session_state.quiz_index
    
    if index >= len(questions):
        score = st.session_state.quiz_score
        total = len(questions)
        st.markdown(f"""
        <div style="text-align: center; padding: 25px;">
            <div style="font-size: 40px;">🎉</div>
            <div style="font-size: 22px; margin: 15px 0;">测验完成！</div>
            <div style="font-size: 32px; color: #667eea; font-weight: bold;">{score} / {total}</div>
            <div style="font-size: 16px; color: #666;">正确率: {int(score/total*100)}%</div>
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📝 再来一次", use_container_width=True):
            st.session_state.quiz_questions = []
            st.rerun()
        return
    
    question = questions[index]
    word = question["word"]
    
    st.progress((index + 1) / len(questions))
    st.caption(f"第 {index + 1} 题 / 共 {len(questions)} 题")
    
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                color: white; padding: 25px; border-radius: 15px; text-align: center; margin: 15px 0;">
        <div style="font-size: 13px; opacity: 0.8;">请选择正确的中文意思</div>
        <div style="font-size: 28px; font-weight: bold; margin: 12px 0;">{word['english']}</div>
        <div style="font-size: 14px; opacity: 0.9;">{word['phonetic']}</div>
    </div>
    """, unsafe_allow_html=True)
    
    if not st.session_state.quiz_answered:
        for i, option in enumerate(question["options"]):
            if st.button(f"{chr(65+i)}. {option}", key=f"opt_{i}", use_container_width=True):
                st.session_state.quiz_answered = True
                if option == question["correct"]:
                    st.session_state.quiz_score += 1
                    mastery = st.session_state.progress.get("mastery", {})
                    mastery[word["id"]] = min(100, mastery.get(word["id"], 0) + 10)
                    st.session_state.progress["mastery"] = mastery
                    save_progress(st.session_state.progress)
                    st.session_state.last_answer_correct = True
                else:
                    mastery = st.session_state.progress.get("mastery", {})
                    mastery[word["id"]] = max(0, mastery.get(word["id"], 0) - 5)
                    st.session_state.progress["mastery"] = mastery
                    save_progress(st.session_state.progress)
                    st.session_state.last_answer_correct = False
                st.rerun()
    else:
        if st.session_state.last_answer_correct:
            st.success("✅ 回答正确！")
        else:
            st.error(f"❌ 回答错误！正确答案: {question['correct']}")
        st.markdown(f"**例句:** {word['example']}")
        if st.button("下一题 →", use_container_width=True):
            st.session_state.quiz_index += 1
            st.session_state.quiz_answered = False
            st.rerun()


def extract_docx_content(file_bytes):
    """深度提取Word文档内容，包括所有展开的章节"""
    from docx import Document
    from docx.oxml.ns import qn
    import io
    
    doc = Document(io.BytesIO(file_bytes))
    all_text = []
    
    # 1. 提取所有段落（包括标题、正文等）
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_text.append(text)
    
    # 2. 提取所有表格内容（深度遍历每个单元格）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 单元格可能包含多个段落
                for para in cell.paragraphs:
                    cell_text = para.text.strip()
                    if cell_text:
                        all_text.append(cell_text)
                # 单元格内可能还有嵌套表格
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for para in nested_cell.paragraphs:
                                if para.text.strip():
                                    all_text.append(para.text.strip())
    
    # 3. 尝试提取文本框和形状中的文本
    try:
        for shape in doc.inline_shapes:
            if hasattr(shape, '_inline'):
                pass
    except:
        pass
    
    # 4. 提取页眉页脚
    try:
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text.strip())
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text.strip())
    except:
        pass
    
    # 5. 尝试从XML中提取所有文本（捕获可能遗漏的内容）
    try:
        from docx.oxml.ns import qn
        body = doc.element.body
        for elem in body.iter():
            if elem.text and elem.text.strip():
                text = elem.text.strip()
                if text not in all_text:
                    all_text.append(text)
            if elem.tail and elem.tail.strip():
                text = elem.tail.strip()
                if text not in all_text:
                    all_text.append(text)
    except:
        pass
    
    return '\n'.join(all_text)


def show_import():
    """导入文档页面"""
    st.markdown("## 📄 导入英文文档")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.rerun()
    
    st.markdown("---")
    st.markdown("""
    上传英文需求文档或技术文档，自动提取其中的单词和句子添加到词库。
    
    支持格式: TXT, MD, DOCX (Word文档)
    
    💡 **提示**: Word文档会自动提取所有层级目录（如2.1, 2.1.1, 2.1.2等）中的内容
    """)
    
    uploaded_file = st.file_uploader("选择文件", type=['txt', 'md', 'docx'])
    
    if uploaded_file:
        # 根据文件类型读取内容
        if uploaded_file.name.endswith('.docx'):
            try:
                file_bytes = uploaded_file.read()
                content = extract_docx_content(file_bytes)
                st.info(f"📄 已从Word文档提取 {len(content)} 字符内容")
            except ImportError:
                st.error("❌ 需要安装 python-docx 库来读取 Word 文档")
                st.code("pip install python-docx", language="bash")
                st.stop()
            except Exception as e:
                st.error(f"❌ 读取 Word 文档失败: {e}")
                st.stop()
        else:
            content = uploaded_file.read().decode('utf-8')
        st.success(f"✅ 文件已上传: {uploaded_file.name}")
        
        # 提取单词
        words = extract_words_from_text(content)
        sentences = extract_sentences_from_text(content)
        
        st.markdown(f"### 📊 提取结果")
        st.markdown(f"- 发现 **{len(words)}** 个英文单词（已去重）")
        st.markdown(f"- 发现 **{len(sentences)}** 个英文句子（已去重）")
        
        # 显示提取的单词（使用文本区域，支持大量数据）
        if words:
            with st.expander(f"📖 查看所有 {len(words)} 个单词", expanded=False):
                # 使用文本区域显示，每行10个单词
                word_lines = []
                for i in range(0, len(words), 10):
                    word_lines.append("  |  ".join(words[i:i+10]))
                st.text_area("单词列表", "\n".join(word_lines), height=400, disabled=True)
                # 提供下载按钮
                st.download_button(
                    "📥 下载单词列表",
                    "\n".join(words),
                    file_name="extracted_words.txt",
                    mime="text/plain"
                )
        
        # 显示提取的句子（使用文本区域）
        if sentences:
            with st.expander(f"💬 查看所有 {len(sentences)} 个句子", expanded=False):
                sentence_text = "\n\n".join([f"{i}. {s}" for i, s in enumerate(sentences, 1)])
                st.text_area("句子列表", sentence_text, height=400, disabled=True)
                # 提供下载按钮
                st.download_button(
                    "📥 下载句子列表",
                    sentence_text,
                    file_name="extracted_sentences.txt",
                    mime="text/plain"
                )
        
        st.markdown("---")
        
        # 添加到词库 - 可选择添加数量
        st.markdown("### ➕ 添加到词库")
        category = st.text_input("词汇分类", value="导入文档")
        scenario = st.text_input("短语场景", value="文档内容")
        
        # 选择添加数量
        col_opt1, col_opt2 = st.columns(2)
        with col_opt1:
            word_limit = st.selectbox("添加单词数量", 
                options=[50, 100, 200, 500, 1000, len(words)],
                format_func=lambda x: f"全部 ({x}个)" if x == len(words) else f"前{x}个",
                index=min(2, len([50, 100, 200, 500, 1000, len(words)]) - 1)
            )
        with col_opt2:
            sentence_limit = st.selectbox("添加句子数量",
                options=[20, 50, 100, len(sentences)],
                format_func=lambda x: f"全部 ({x}个)" if x == len(sentences) else f"前{x}个",
                index=min(1, len([20, 50, 100, len(sentences)]) - 1)
            )
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📖 添加单词到词库", use_container_width=True):
                custom = load_custom_data()
                existing_words = set(w["english"].lower() for w in OTA_VOCABULARY + custom.get("vocabulary", []))
                
                new_words = []
                for i, word in enumerate(words[:word_limit]):
                    if word.lower() not in existing_words:
                        existing_words.add(word.lower())  # 防止重复添加
                        new_words.append({
                            "id": f"custom_v{len(custom.get('vocabulary', [])) + len(new_words) + 1}",
                            "english": word,
                            "chinese": "(待翻译)",
                            "phonetic": "",
                            "example": "",
                            "category": category
                        })
                
                if new_words:
                    custom["vocabulary"] = custom.get("vocabulary", []) + new_words
                    if save_custom_data(custom):
                        st.success(f"✅ 已添加 {len(new_words)} 个新单词到词库！")
                        st.info(f"📚 当前词库共有 {len(OTA_VOCABULARY) + len(custom['vocabulary'])} 个单词")
                else:
                    st.info("没有新单词需要添加（可能都已存在）")
        
        with col2:
            if st.button("💬 添加句子到词库", use_container_width=True):
                custom = load_custom_data()
                existing_phrases = set(p["english"].lower() for p in OTA_PHRASES + custom.get("phrases", []))
                
                new_phrases = []
                for i, sentence in enumerate(sentences[:sentence_limit]):
                    if sentence.lower() not in existing_phrases:
                        existing_phrases.add(sentence.lower())
                        new_phrases.append({
                            "id": f"custom_p{len(custom.get('phrases', [])) + len(new_phrases) + 1}",
                            "english": sentence,
                            "chinese": "(待翻译)",
                            "scenario": scenario
                        })
                
                if new_phrases:
                    custom["phrases"] = custom.get("phrases", []) + new_phrases
                    if save_custom_data(custom):
                        st.success(f"✅ 已添加 {len(new_phrases)} 个句子到词库！")
                        st.info(f"💬 当前词库共有 {len(OTA_PHRASES) + len(custom['phrases'])} 个短语")
                else:
                    st.info("没有新句子需要添加（可能都已存在）")
        
        # 一键添加全部
        st.markdown("---")
        if st.button("🚀 一键添加全部单词和句子", use_container_width=True, type="primary"):
            custom = load_custom_data()
            
            # 添加单词
            existing_words = set(w["english"].lower() for w in OTA_VOCABULARY + custom.get("vocabulary", []))
            new_words = []
            for word in words:
                if word.lower() not in existing_words:
                    existing_words.add(word.lower())
                    new_words.append({
                        "id": f"custom_v{len(custom.get('vocabulary', [])) + len(new_words) + 1}",
                        "english": word,
                        "chinese": "(待翻译)",
                        "phonetic": "",
                        "example": "",
                        "category": category
                    })
            
            # 添加句子
            existing_phrases = set(p["english"].lower() for p in OTA_PHRASES + custom.get("phrases", []))
            new_phrases = []
            for sentence in sentences:
                if sentence.lower() not in existing_phrases:
                    existing_phrases.add(sentence.lower())
                    new_phrases.append({
                        "id": f"custom_p{len(custom.get('phrases', [])) + len(new_phrases) + 1}",
                        "english": sentence,
                        "chinese": "(待翻译)",
                        "scenario": scenario
                    })
            
            custom["vocabulary"] = custom.get("vocabulary", []) + new_words
            custom["phrases"] = custom.get("phrases", []) + new_phrases
            
            if save_custom_data(custom):
                st.success(f"✅ 已添加 {len(new_words)} 个单词 + {len(new_phrases)} 个句子！")
                st.info(f"📚 词库总计: {len(OTA_VOCABULARY) + len(custom['vocabulary'])} 单词, {len(OTA_PHRASES) + len(custom['phrases'])} 短语")


def show_manage():
    """管理词库"""
    st.markdown("## ⚙️ 管理词库")
    if st.button("← 返回首页"):
        st.session_state.current_page = "首页"
        st.rerun()
    
    st.markdown("---")
    
    custom = load_custom_data()
    custom_vocab = custom.get("vocabulary", [])
    custom_phrases = custom.get("phrases", [])
    
    st.markdown(f"### 📊 词库统计")
    col1, col2 = st.columns(2)
    with col1:
        st.metric("内置词汇", len(OTA_VOCABULARY))
        st.metric("内置短语", len(OTA_PHRASES))
    with col2:
        st.metric("自定义词汇", len(custom_vocab))
        st.metric("自定义短语", len(custom_phrases))
    
    st.markdown("---")
    
    # 手动添加词汇
    st.markdown("### ➕ 手动添加词汇")
    with st.expander("添加新词汇"):
        new_english = st.text_input("英文单词")
        new_chinese = st.text_input("中文翻译")
        new_phonetic = st.text_input("音标 (可选)")
        new_example = st.text_input("例句 (可选)")
        new_category = st.text_input("分类", value="自定义")
        
        if st.button("添加词汇"):
            if new_english and new_chinese:
                custom["vocabulary"] = custom.get("vocabulary", []) + [{
                    "id": f"custom_v{len(custom.get('vocabulary', [])) + 1}",
                    "english": new_english,
                    "chinese": new_chinese,
                    "phonetic": new_phonetic,
                    "example": new_example,
                    "category": new_category
                }]
                save_custom_data(custom)
                st.success("✅ 词汇已添加！")
                st.rerun()
            else:
                st.error("请填写英文和中文")
    
    # 手动添加短语
    with st.expander("添加新短语"):
        new_phrase_en = st.text_input("英文短语/句子")
        new_phrase_cn = st.text_input("中文翻译 ")
        new_scenario = st.text_input("场景", value="自定义")
        
        if st.button("添加短语"):
            if new_phrase_en and new_phrase_cn:
                custom["phrases"] = custom.get("phrases", []) + [{
                    "id": f"custom_p{len(custom.get('phrases', [])) + 1}",
                    "english": new_phrase_en,
                    "chinese": new_phrase_cn,
                    "scenario": new_scenario
                }]
                save_custom_data(custom)
                st.success("✅ 短语已添加！")
                st.rerun()
            else:
                st.error("请填写英文和中文")
    
    st.markdown("---")
    
    # 查看自定义词汇
    if custom_vocab:
        st.markdown("### 📖 自定义词汇列表")
        for word in custom_vocab[:20]:
            st.markdown(f"- **{word['english']}** - {word['chinese']}")
    
    # 清空自定义数据
    st.markdown("---")
    if st.button("🗑️ 清空所有自定义数据", type="secondary"):
        save_custom_data({"vocabulary": [], "phrases": []})
        st.success("已清空自定义数据")
        st.rerun()


# ==================== 主程序 ====================
def main():
    with st.sidebar:
        st.markdown("## 📚 OTA英语 V2.0")
        st.markdown("---")
        if st.button("🏠 首页", use_container_width=True):
            st.session_state.current_page = "首页"
            st.rerun()
        if st.button("📖 词汇", use_container_width=True):
            st.session_state.current_page = "词汇"
            st.rerun()
        if st.button("💬 短语", use_container_width=True):
            st.session_state.current_page = "短语"
            st.rerun()
        if st.button("🎴 闪卡", use_container_width=True):
            st.session_state.current_page = "闪卡"
            st.rerun()
        if st.button("📝 测验", use_container_width=True):
            st.session_state.current_page = "测验"
            st.rerun()
        if st.button("📄 导入", use_container_width=True):
            st.session_state.current_page = "导入"
            st.rerun()
        if st.button("⚙️ 管理", use_container_width=True):
            st.session_state.current_page = "管理"
            st.rerun()
        
        st.markdown("---")
        st.markdown("### 📱 添加到主屏幕")
        st.markdown("Safari → 分享 → 添加到主屏幕")
    
    page = st.session_state.current_page
    if page == "首页":
        show_home()
    elif page == "词汇":
        show_vocabulary()
    elif page == "短语":
        show_phrases()
    elif page == "闪卡":
        show_flashcards()
    elif page == "测验":
        show_quiz()
    elif page == "导入":
        show_import()
    elif page == "管理":
        show_manage()

if __name__ == "__main__":
    main()

